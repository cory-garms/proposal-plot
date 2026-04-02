"""
Convert a draft's markdown-style content to a .docx file in memory.

Handles the heading/bullet structure produced by the RAG generator:
  ## Section Title  -> Heading 1
  ### Sub-title     -> Heading 2
  - bullet item     -> List Bullet style
  blank line        -> paragraph break
  everything else   -> Normal body paragraph
"""
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _parse_lines(content: str) -> list[tuple[str, str]]:
    """
    Returns list of (type, text) tuples.
    Types: 'h1', 'h2', 'bullet', 'body', 'empty'
    """
    tokens = []
    for line in content.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("## "):
            tokens.append(("h1", stripped[3:].strip()))
        elif stripped.startswith("### "):
            tokens.append(("h2", stripped[4:].strip()))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            tokens.append(("bullet", stripped[2:].strip()))
        elif stripped == "":
            tokens.append(("empty", ""))
        else:
            tokens.append(("body", stripped))
    return tokens


def build_docx(title: str, agency: str, content: str) -> bytes:
    """
    Build a .docx document and return its bytes.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Cover header
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph(f"Agency: {agency}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = None  # default gray via style

    doc.add_paragraph("")  # spacer

    tokens = _parse_lines(content)
    # Collapse consecutive 'empty' tokens
    prev_empty = False
    for kind, text in tokens:
        if kind == "empty":
            if not prev_empty:
                doc.add_paragraph("")
            prev_empty = True
            continue
        prev_empty = False

        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
